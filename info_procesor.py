import numpy as np
import pandas as pd 
import datetime as dt
import json
from docx import Document
from docx.shared import Inches
from pandas import Series, DataFrame
import graphic_creator
import getPython
import quickstart

def get_doc(tablas):
    now = dt.datetime.now()
    document = Document()

    # Todas las variables deseadas se encuentran en este arreglo.
    arr = ['corriente_x', 'corriente_y', 'corriente_z', 'potencia_x', 'potencia_y', 'potencia_z']


    # Abre logs y consigue el numero del reporte.
    with open('LOGS.json') as js:
        data = json.load(js)
        num_reporte = data['reporte']
        data['fecha'] = str(now)
        data['reporte'] = data['reporte'] + 1

    #Actualiza Logs.
    with open('LOGS.json', 'w') as j:
        json.dump(data, j)

    r_str = 'Reporte ' + str(num_reporte)

    jsonf = getPython.get_json()
    dataf = DataFrame(jsonf['rows']).fillna(0) # fillna(x) --> transforma todos los valores de na (nulos) a el valor x.
 
    # Esta parte crea las graficas, se puee resumir a un ciclo "for" como se desee que el reporte quede.
    graphic_creator.creador_graficas(dataf, 0, 'corriente_x', 'corriente_y')
    graphic_creator.creador_graficas(dataf, 1, 'corriente_y', 'corriente_z')
    graphic_creator.creador_graficas(dataf, 2, 'corriente_x', 'corriente_z')
    graphic_creator.creador_graficas(dataf, 3, 'potencia_x', 'potencia_y')
    graphic_creator.creador_graficas(dataf, 4, 'potencia_y', 'potencia_z')
    graphic_creator.creador_graficas(dataf, 5, 'potencia_x', 'potencia_z')


    # ("x" y "p") son nombres de partes del documento.
    # Para todo el manejo del documento (https://python-docx.readthedocs.io/en/latest/user/quickstart.html)
    x = document.add_heading(str('Reporte ' + str(num_reporte)), 0)
    p = document.add_paragraph(r_str + ' en la fecha ' + str(now))

    for x in range(0,5):
        document.add_picture(str('grafica ' + str(x) + '.png'), width=Inches(4.5))
    
    for img in arr:
        document.add_heading('INFORMACION SOBRE ' + img)
        str1 = graphic_creator.get_promedio(dataf, img)
        str2 = graphic_creator.get_media(dataf, img)
        str3 = graphic_creator.get_moda(dataf, img)
        str4 = graphic_creator.get_pico(dataf, img)    
        p1 = document.add_paragraph(str1)
        p2 = document.add_paragraph(str2)
        p3 = document.add_paragraph(str3)
        p4 = document.add_paragraph(str4)

    if tablas == 1:
        print(document.add_heading('TABLAS'))

    nombre_reporte = r_str + '.docx'
    document.save(nombre_reporte)

    link = quickstart.subir_archivo(nombre_reporte)

    return link